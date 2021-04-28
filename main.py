import docx
import sys
from app_config import AppConfig
from terms_loader import TermsLoader


def doc_export(terms: list, lang: str):
    for i in range(len(terms)):
        terms[i] = terms[i].split(" - ", 1)
    doc = docx.Document()
    doc.add_heading("Глоссарий", 0)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    if lang.lower() == "рус":
        header_cells[0].text = "Термин"
        header_cells[1].text = "Значение"
    else:
        header_cells[0].text = "Термин"
        header_cells[1].text = "Терминнің мағынасы"
    i = 0
    term = ""
    description = ""
    try:
        for term, description in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = term
            row_cells[1].text = description
            i += 1
    except:
        print(f"Exception occurred: {sys.exc_info()[0]}")
        print(f"Term: {term} Description: {description} I: {i}")
        print(terms[i])
    if lang.lower() == "рус":
        filename = "glossary_ru.docx"
    else:
        filename = "glossary_kz.docx"
    doc.save(filename)


def main():
    AppConfig.terms_count = -1
    AppConfig.lang = "-"
    while AppConfig.terms_count == -1:
        AppConfig.terms_count = int(input("Колличество терминов: "))
        if AppConfig.terms_count >= 502 or AppConfig.terms_count < 1:
            print("Таблица содержит 502 термина и колиичество терминов не может быть меньше 1")
            AppConfig.terms_count = -1
    while AppConfig.lang == "-":
        AppConfig.lang = input("Язык(РУС/каз): ")
        if AppConfig.lang.strip() == "":
            AppConfig.lang = "рус"
        elif AppConfig.lang.lower() != "рус" and AppConfig.lang.lower() != "каз":
            AppConfig.lang = "-"
    selected_terms = TermsLoader.get_terms()
    export_terms_to_txt(selected_terms, AppConfig.lang)
    doc_export(selected_terms, AppConfig.lang)


def export_terms_to_txt(selected_terms: list, lang: str):
    if lang.lower() == "рус":
        txt_filename = "glossary_ru.txt"
    else:
        txt_filename = "glossary_kz.txt"
    with open(txt_filename, "w", encoding="UTF-8") as glossaryFile:
        for term in selected_terms:
            glossaryFile.write(term)


if __name__ == "__main__":
    main()
