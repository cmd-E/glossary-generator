import requests
from bs4 import BeautifulSoup
import random
import docx
import sys


def parse_terms(raw_terms: list, lang: str) -> list:
    description = ""
    term = ""
    terms_list = []
    j = 0
    for tr in raw_terms:
        if j % 2 != 0:
            if lang.lower() == "рус":
                description = tr.find("td").find("p").getText()
                description = description.replace("\r", "")
                description = description.replace("\n", "")
            else:
                description = tr.find_all("td")[1].find("p").getText()
                description = description.replace("\r", "")
                description = description.replace("\n", "")
        else:
            if lang.lower() == "рус":
                term = tr.find("td").find("p").getText()
                term = term.replace("\r", "")
                term = term.replace("\n", "")
            else:
                term = tr.find_all("td")[1].find("p").getText()
                term = term.replace("\r", "")
                term = term.replace("\n", "")
        if term == "" or description == "":
            j += 1
            continue
        else:
            terms_list.append([
                term, description
            ])
            term = ""
            description = ""
            j += 1
    return terms_list


def doc_export(terms: list, lang: str):
    for i in range(len(terms)):
        terms[i] = terms[i].split(" - ", 1)
    doc = docx.Document()
    doc.add_heading("Глоссарий", 0)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    if lang.lower() == "рус":
        header_cells[0].text = "Термин"
        header_cells[1].text = "Значение"
    else:
        header_cells[0].text = "Термин"
        header_cells[1].text = "Терминнің мағынасы"
    i = 0
    term = ""
    description = ""
    try:
        for term, description in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = term
            row_cells[1].text = description
            i += 1
    except:
        print(f"Exception occurred: {sys.exc_info()[0]}")
        print(f"Term: {term} Description: {description} I: {i}")
        print(terms[i])
    if lang.lower() == "рус":
        filename = "glossary_ru.docx"
    else:
        filename = "glossary_kz.docx"
    doc.save(filename)


def get_random_terms(terms_list: list, terms_count: int) -> list:
    used_terms = []
    selected_terms = []
    for i in range(terms_count):
        random_term = terms_list[random.randint(0, len(terms_list) - 1)]
        formatted_term = f"{random_term[0]} - {random_term[1]}\n"
        if formatted_term not in used_terms:
            selected_terms.append(formatted_term)
            used_terms.append(formatted_term)
        else:
            terms_count += 1
            continue
    return selected_terms


def main():
    terms_count = -1
    lang = "-"
    while terms_count == -1:
        terms_count = int(input("Колличество терминов: "))
        if terms_count >= 502 or terms_count < 1:
            print("Таблица содержит 502 термина и колиичество терминов не может быть меньше 1")
            terms_count = -1
    while lang == "-":
        lang = input("Язык(РУС/каз): ")
        if lang.strip() == "":
            lang = "рус"
        elif lang.lower() != "рус" and lang.lower() != "каз":
            lang = "-"
    raw_terms = load_raw_terms()
    terms_list = parse_terms(raw_terms, lang)
    selected_terms = get_random_terms(terms_list, terms_count)
    export_terms_to_txt(selected_terms, lang)
    doc_export(selected_terms, lang)


def load_raw_terms() -> list:
    url = "http://libr.aues.kz/facultet/frts/kaf_aes/52/umm/aes_1.htm"
    response = ""
    try:
        response = requests.get(url)
    except:
        print("Unexpected error occurred while getting response from glossary page: ",
              sys.exc_info()[0])
        exit()

    soup = ""
    try:
        soup = BeautifulSoup(response.content, "html.parser")
    except:
        print("Unexpected error occurred while parsing response from glossary page: ",
              sys.exc_info()[0])

    trs = soup.find("table").find_all("tr")
    trs = trs[1:]
    return trs


def export_terms_to_txt(selected_terms: list, lang: str):
    if lang.lower() == "рус":
        txt_filename = "glossary_ru.txt"
    else:
        txt_filename = "glossary_kz.txt"
    with open(txt_filename, "w", encoding="UTF-8") as glossaryFile:
        for term in selected_terms:
            glossaryFile.write(term)


if __name__ == "__main__":
    main()
