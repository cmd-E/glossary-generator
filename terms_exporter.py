import docx


class TermsExporter:
    @classmethod
    def doc_export(cls, terms: list, lang: str):
        for i in range(len(terms)):
            terms[i] = terms[i].split(" - ", 1)
        doc = docx.Document()
        doc.add_heading("Глоссарий", 0)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        if lang.lower() == "рус":
            header_cells[0].text = "Термин"
            header_cells[1].text = "Значение"
        else:
            header_cells[0].text = "Термин"
            header_cells[1].text = "Терминнің мағынасы"
        i = 0
        term = ""
        description = ""
        try:
            for term, description in terms:
                row_cells = table.add_row().cells
                row_cells[0].text = term
                row_cells[1].text = description
                i += 1
        except:
            print(f"Exception occurred: {sys.exc_info()[0]}")
            print(f"Term: {term} Description: {description} I: {i}")
            print(terms[i])
        if lang.lower() == "рус":
            filename = "glossary_ru.docx"
        else:
            filename = "glossary_kz.docx"
        doc.save(filename)

    @classmethod
    def export_terms_to_txt(cls, selected_terms: list, lang: str):
        if lang.lower() == "рус":
            txt_filename = "glossary_ru.txt"
        else:
            txt_filename = "glossary_kz.txt"
        with open(txt_filename, "w", encoding="UTF-8") as glossaryFile:
            for term in selected_terms:
                glossaryFile.write(term)
